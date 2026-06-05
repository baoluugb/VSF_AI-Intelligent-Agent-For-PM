"""Export utilities for VSF Project Intelligence Agent reports.

This module converts:
1. Markdown reports (output/report.md) -> styled Word documents (output/report.docx)
2. JSON concerns list (output/concerns.json) -> Excel spreadsheets (output/concerns.xlsx)
   with conditional formatting for Severity = 5 items.
"""
from __future__ import annotations

import os
import re
import json
import logging
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

logger = logging.getLogger(__name__)


def add_formatted_runs(paragraph, text: str) -> None:
    """Helper to parse **bold** and _italic_ text and append them to a paragraph."""
    # Pattern split for **bold** or _italic_
    pattern = re.compile(r'(\*\*[^*]+\*\*|_[^_]+_)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('_') and part.endswith('_'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def export_report_to_docx(markdown_path: str, docx_path: str) -> None:
    """Parse report.md into a beautifully styled Word Document (.docx)."""
    if not os.path.exists(markdown_path):
        raise FileNotFoundError(f"Markdown report not found at: {markdown_path}")

    doc = Document()

    # Define Document Typography & Styles
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x2B, 0x2D, 0x42) # Slate Gray
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Customize Headings
    heading_styles = {
        'Heading 1': (16, RGBColor(0x1D, 0x35, 0x57), True), # Deep Navy
        'Heading 2': (14, RGBColor(0x45, 0x7B, 0x9D), True), # Muted Blue
        'Heading 3': (12, RGBColor(0x1D, 0x35, 0x57), True), # Deep Navy Bold
    }

    for name, (size, color, bold) in heading_styles.items():
        if name in doc.styles:
            h_style = doc.styles[name]
            h_font = h_style.font
            h_font.name = 'Calibri'
            h_font.size = Pt(size)
            h_font.color.rgb = color
            h_font.bold = bold
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)
            h_style.paragraph_format.keep_with_next = True

    # Parse and write
    with open(markdown_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False

    for line in lines:
        stripped = line.strip()

        # Code blocks toggle
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line.rstrip('\n'))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x6A, 0x73, 0x7D)
            continue

        # Headings
        if stripped.startswith('#'):
            depth = len(stripped) - len(stripped.lstrip('#'))
            title_text = stripped.lstrip('#').strip()
            if title_text:
                doc.add_heading(title_text, level=min(depth, 3))
            continue

        # Bullet list items
        if stripped.startswith('- ') or stripped.startswith('* '):
            list_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_runs(p, list_text)
            continue

        # Empty lines
        if not stripped:
            continue

        # Standard paragraph
        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)

    doc.save(docx_path)
    logger.info("Successfully exported %s to %s", markdown_path, docx_path)


def export_concerns_to_excel(json_path: str, excel_path: str) -> None:
    """Read concerns.json, structure as a flat table, and export to Excel (.xlsx)

    Highlights rows with Severity = 5 with a premium light red color scheme.
    """
    if not os.path.exists(json_path):
        raise FileNotFoundError(f"JSON concerns file not found at: {json_path}")

    with open(json_path, 'r', encoding='utf-8') as f:
        concerns = json.load(f)

    flat_concerns = []
    for c in concerns:
        details = c.get("details", {})
        
        # Pull specific detail fields depending on the issue type
        days_val = details.get("days_stalled") or details.get("days_remaining") or details.get("days_open")
        
        row = {
            "Type": c.get("type", ""),
            "Task ID": c.get("task_id", ""),
            "Severity": c.get("severity", 1),
            "Explanation": c.get("explanation", ""),
            "Assignee": c.get("assignee", "") or "Unassigned",
            "Source IDs": ", ".join(c.get("source_ids", [])) if isinstance(c.get("source_ids"), list) else str(c.get("source_ids") or ""),
            "Status": details.get("status", ""),
            "Due Date": details.get("due_date", ""),
            "Days Stalled/Remaining/Open": days_val if days_val is not None else "",
            "Blocked Dependency Count": details.get("dependent_count", ""),
            "Evidence Note": details.get("note_id", ""),
            "Evidence Content": details.get("evidence", "")
        }
        flat_concerns.append(row)

    df = pd.DataFrame(flat_concerns)

    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Concerns & Risks')

        workbook = writer.book
        worksheet = writer.sheets['Concerns & Risks']

        # Theme Colors: Deep Slate Blue Header, Light Red High Alert
        header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
        header_font = Font(name="Segoe UI", size=11, bold=True, color="FFFFFF")

        # Soft Red background (#FFD8D8) + Dark Red bold font (#9C0006)
        sev5_fill = PatternFill(start_color="FFD8D8", end_color="FFD8D8", fill_type="solid")
        sev5_font = Font(name="Segoe UI", size=10, bold=True, color="9C0006")

        normal_font = Font(name="Segoe UI", size=10)
        thin_border = Border(
            left=Side(style='thin', color='E0E0E0'),
            right=Side(style='thin', color='E0E0E0'),
            top=Side(style='thin', color='E0E0E0'),
            bottom=Side(style='thin', color='E0E0E0')
        )

        # Show Excel grid lines
        try:
            worksheet.sheet_view.showGridLines = True
        except AttributeError:
            try:
                worksheet.views.sheetView[0].showGridLines = True
            except Exception:
                pass

        # Format header row
        worksheet.row_dimensions[1].height = 28
        for col_num in range(1, len(df.columns) + 1):
            cell = worksheet.cell(row=1, column=col_num)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=False)
            cell.border = thin_border

        # Iterate and apply styles & conditional formats
        for row_idx in range(2, len(df) + 2):
            severity_val = df.iloc[row_idx - 2]["Severity"]
            is_sev5 = (severity_val == 5 or str(severity_val) == '5')

            worksheet.row_dimensions[row_idx].height = 20

            for col_idx in range(1, len(df.columns) + 1):
                cell = worksheet.cell(row=row_idx, column=col_idx)
                cell.border = thin_border
                cell.font = normal_font

                # Clean cell alignments based on content type
                if col_idx in [2, 3, 7, 8, 9, 10]:  # Center code references, numbers & statuses
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                else:
                    cell.alignment = Alignment(horizontal="left", vertical="center")

                # Highlight row if severity is 5
                if is_sev5:
                    cell.fill = sev5_fill
                    cell.font = sev5_font

        # Auto-adjust column widths with safety limits
        for col in worksheet.columns:
            max_len = 0
            col_letter = get_column_letter(col[0].column)

            for cell in col:
                val_str = str(cell.value or '')
                if cell.row == 1:
                    max_len = max(max_len, len(val_str) + 4)
                else:
                    # Truncate long texts for column sizing limits
                    max_len = max(max_len, min(len(val_str), 35))

            worksheet.column_dimensions[col_letter].width = max(max_len + 3, 12)

    logger.info("Successfully exported %s to %s", json_path, excel_path)
