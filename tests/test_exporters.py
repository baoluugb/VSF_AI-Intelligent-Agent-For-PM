import os
import json
import pytest
from docx import Document
from openpyxl import load_workbook
from src.exporters import export_report_to_docx, export_concerns_to_excel


def test_export_report_to_docx(tmp_path):
    # Create a dummy markdown file
    md_content = """# Title 1
## Title 2
Some normal paragraph text with **bold text** and _italic text_.

- First bullet item with **bold context**
- Second bullet item

```
code line 1
code line 2
```
"""
    md_file = tmp_path / "report.md"
    md_file.write_text(md_content, encoding="utf-8")

    docx_file = tmp_path / "report.docx"
    export_report_to_docx(str(md_file), str(docx_file))

    # Assert docx file is created
    assert os.path.exists(docx_file)
    assert os.path.getsize(docx_file) > 0

    # Check document content
    doc = Document(str(docx_file))
    paragraphs = [p.text for p in doc.paragraphs]

    assert "Title 1" in paragraphs
    assert "Title 2" in paragraphs
    assert "Some normal paragraph text with bold text and italic text." in paragraphs
    assert "First bullet item with bold context" in paragraphs
    assert "code line 1" in paragraphs

    # Check formatting runs in bullet list
    bullet_p = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
    assert len(bullet_p) == 2
    runs = bullet_p[0].runs
    assert len(runs) >= 2

    # Check if bold run exists
    bold_runs = [r for r in runs if r.bold]
    assert len(bold_runs) == 1
    assert bold_runs[0].text == "bold context"


def test_export_concerns_to_excel(tmp_path):
    # Create a dummy concerns.json file
    concerns_data = [
        {
            "type": "deadline_risk",
            "task_id": "FLINK-40",
            "severity": 5,
            "explanation": "Deadline đã quá hạn 2 ngày, status vẫn 'Reopened'.",
            "assignee": "Jack Jackson",
            "source_ids": ["FLINK-40"],
            "details": {
                "days_remaining": -2,
                "due_date": "2025-05-28",
                "status": "Reopened"
            }
        },
        {
            "type": "stalled_task",
            "task_id": "CASSANDRA-96",
            "severity": 3,
            "explanation": "No updates in 110 days.",
            "assignee": "Duc Anh",
            "source_ids": ["CASSANDRA-96"],
            "details": {
                "days_stalled": 110,
                "status": "In Progress"
            }
        }
    ]

    json_file = tmp_path / "concerns.json"
    with open(json_file, "w", encoding="utf-8") as f:
        json.dump(concerns_data, f)

    excel_file = tmp_path / "concerns.xlsx"
    export_concerns_to_excel(str(json_file), str(excel_file))

    assert os.path.exists(excel_file)
    assert os.path.getsize(excel_file) > 0

    # Read sheet and check content
    wb = load_workbook(str(excel_file))
    sheet = wb['Concerns & Risks']

    # Check headers
    headers = [cell.value for cell in sheet[1]]
    assert "Task ID" in headers
    assert "Severity" in headers

    # Row 2 is FLINK-40, Row 3 is CASSANDRA-96
    assert sheet.cell(row=2, column=2).value == "FLINK-40"
    assert sheet.cell(row=2, column=3).value == 5

    assert sheet.cell(row=3, column=2).value == "CASSANDRA-96"
    assert sheet.cell(row=3, column=3).value == 3

    # Check styling for severity = 5 row (row 2)
    flink_cell = sheet.cell(row=2, column=2)
    cassandra_cell = sheet.cell(row=3, column=2)

    # Check that fill was applied to flink cell
    assert flink_cell.fill.start_color.rgb in ["00FFD8D8", "FFFFD8D8", "FFD8D8"]
    # Check that fill was not applied to cassandra cell
    assert cassandra_cell.fill.fill_type is None
